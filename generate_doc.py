from docx import Document

doc = Document()
p = doc.add_paragraph()
run1 = p.add_run('S')
run1.font.italic = True
run1.font.name = 'Cambria Math'
run2 = p.add_run(' = 1 - max(')
run2.font.name = 'Cambria Math'
run3 = p.add_run('p')
run3.font.italic = True
run3.font.name = 'Cambria Math'
run4 = p.add_run('out')
run4.font.subscript = True
run4.font.name = 'Cambria Math'
run5 = p.add_run(')')
run5.font.name = 'Cambria Math'

doc.save('d:\\ttk-S9-260430-v3.4\\formula.docx')
